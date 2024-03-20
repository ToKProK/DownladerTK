from selenium import webdriver
import time
from fake_useragent import UserAgent
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
import os
from docx import Document
from glob import glob

my_list = []
path_list = []
count_chapters = 0
nomer_glavi = 1
title_name = ""
arthist = ""
time_pause = 2
def parse_ranobe(link, path=None):
    # try:
        global nomer_glavi
        nomer_glavi = 1
        #Настройка юзер агента
        user = UserAgent()
        user_now = user.random
        options_Chr = webdriver.ChromeOptions()
        options_Fire = webdriver.FirefoxOptions()
        #Делаем браузер невимым
        # options_Chr.add_argument("--headless=new")
        options_Chr.add_argument(f"user-agent={user_now}")

        # options_Fire.add_argument("-headless")#https://ru.stackoverflow.com/questions/1330358/%D0%9D%D0%B5-%D1%80%D0%B0%D0%B1%D0%BE%D1%82%D0%B0%D0%B5%D1%82-headless-firefox-selenium
        options_Fire.add_argument(f"user-agent={user_now}")
        #Получение драйвера и вставка юзер агента
        try:
            driver = webdriver.Chrome(options=options_Chr)#Попытка запустить драйвер chrome
        except:
            driver = webdriver.Firefox(options=options_Fire)#Если не получилость пытаемся запустить драйвер firefox

        #Растягиваем окно во всю ширину.
        driver.maximize_window()
        #Переход по ссылке
        driver.get(link)
        time.sleep(time_pause)

        #Попытка получить название
        title_name = driver.find_element(By.CLASS_NAME, "media-name__main").text

        
        #Нажимаю на кнопку начать читать
        driver.find_element(By.LINK_TEXT, "Начать читать").click()
        time.sleep(1)
        #Кликаем по вкладке глав
        driver.find_element(By.XPATH, "//div[@data-reader-modal='chapters']").click()
        chapters = driver.find_elements(By.CLASS_NAME, "menu__item")
        #Получаем количество глав
        count_chapters = (len(chapters))
        print(f"Количесвто глав = {count_chapters}")
        driver.find_element(By.CLASS_NAME, "popup_side").click()


        #Начало работы с документом
        document = Document()
        document.add_heading(f"{title_name}", 0)
        for i in range(0, count_chapters):
            #Получаю название главы
            driver.find_element(By.XPATH, "//div[@data-reader-modal='chapters']").click()
            time.sleep(0.2)
            name_chapter = driver.find_element(By.CLASS_NAME, "menu__item_active").text
            driver.find_element(By.CLASS_NAME, "popup_side").click()
            #Пишу в документ название главы
            document.add_heading(f"{name_chapter}")
            #Получаю весь текст главы
            text_title = driver.find_element(By.CLASS_NAME, "reader-container").find_elements(By.XPATH, "//p")
            
            #Перебираю весь найденный текс и удаляю пустые строки (без понятия откуда они беруться ┐(￣ヘ￣;)┌ )
            corect_text = list()
            for a in text_title:
                 if a.text != "":
                      corect_text.append(a.text)
            for paragraph in corect_text:
                 document.add_paragraph(paragraph)
            time.sleep(0.4)
            try:
                driver.find_element(By.LINK_TEXT, "Следующая глава").click()
            except:
                ""
            text_title.clear()
        title_name = Check_file_name(title_name)
        create_directory(path=path, title_name=title_name)
        if path != 'None':
            document.save(f"{path}/{title_name}.docx")
        else:
            document.save(f"{title_name}.docx") 
        print("Все главы скачанны")
        return True
    # except Exception as ex:
    #     print(ex)
    #     return False
    # finally:
    #     driver.close()
    #     driver.quit()

def create_directory(path, title_name):
    if path != None or path != "": # Тут создаётся папка по выбранному пользователем пути, куда бует все складироваться
        Full_path_dir = os.path.join(path, title_name)# Путь для создания папки, в выбранном каталоге
    else: # Если пользователь не указал путь
        work_dir = os.getcwd() # Узнаём  путь проекта
        Full_path_dir = os.path.join(work_dir, title_name)# Путь для создания папки, в каталоге проекта
    
    if os.path.isdir(Full_path_dir) != True:
        os.mkdir(Full_path_dir)# Создаём директорию в которй будем работать
        print("Папка созданна")
    else:
        print("Папка уже была")
    
    return Full_path_dir

def Del_image(path): # Удаляем все скачанные картинки в папке, после конвертации их в пдф
    pattern = '*.png'
    pattern_path = os.path.join(path, pattern)# Шаблон по которому будут удалятся картинки
    image_path = glob(pattern_path)#Использую библиотеку glob
    for file in image_path:
        os.remove(file)
def Check_file_name(file_name): # Замена символов, которые не могут быть в название файлов
    char_remov = ['*', '|', "\\", ':', '"', '<', '>', '?','/' ]
    for char in char_remov:
        file_name = file_name.replace(char, "#")
    return file_name

def main():
    link = "https://ranobelib.me/about-the-reckless-girl-who-kept-challenging-a-reborn-man-like-me?section=info"
    parse_ranobe(link=link)

if __name__=="__main__":
    main()